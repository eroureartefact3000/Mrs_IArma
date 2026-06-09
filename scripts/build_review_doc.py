"""Generate a Word document for the DA to review the rationale outputs.

Each board occupies its own page(s) with: heading, embedded image, one-liner,
4-axis scores, weighted total, jury verdict, per-axis rationale, creative
mechanisms, and a feedback placeholder.

Workflow:
    1. uv run python build_review_doc.py --input data/<file>.jsonl --output data/<file>.docx --title "..."
    2. Upload .docx to Google Drive
    3. Right-click → 'Open with Google Docs' (auto-converts)
    4. Share with the DA (Comment access is enough)
    5. DA highlights text and adds comments via Insert > Comment (Cmd+Opt+M)
"""
import argparse
import io
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

import pillow_avif  # noqa: F401  # register AVIF support

ROOT = Path(__file__).resolve().parent.parent

TIER_ORDER = {
    "Grand Prix": 0,
    "Gold": 1,
    "Silver": 2,
    "Bronze": 3,
    "Shortlist": 4,
    "Loser": 5,
}
TIER_COLOR = {
    "Grand Prix": RGBColor(0x1A, 0x1A, 0x1A),
    "Gold": RGBColor(0xB8, 0x86, 0x0B),
    "Silver": RGBColor(0x70, 0x70, 0x70),
    "Bronze": RGBColor(0xB8, 0x73, 0x33),
    "Shortlist": RGBColor(0x55, 0x55, 0x55),
    "Loser": RGBColor(0xA8, 0x30, 0x2A),
}
CONSISTENCY_COLOR = {
    "expected": RGBColor(0x2D, 0x6E, 0x3E),
    "above": RGBColor(0xB5, 0x81, 0x05),
    "below": RGBColor(0xA8, 0x30, 0x2A),
    "n/a": RGBColor(0x88, 0x88, 0x88),
}


def resize_for_doc(src: Path, max_width: int = 1300, quality: int = 82) -> io.BytesIO:
    img = Image.open(src).convert("RGB")
    if img.width > max_width:
        ratio = max_width / img.width
        img = img.resize((max_width, int(img.height * ratio)), Image.Resampling.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=quality, optimize=True)
    buf.seek(0)
    return buf


def add_label(paragraph, text: str, *, color: RGBColor | None = None, size: int = 9):
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def render_board(doc: Document, r: dict) -> None:
    w = r["why_it_won"]
    tier = r["tier"]

    # Section heading: tier + campaign + brand
    campaign = r["extracted"].get("campaign") or "(no campaign name)"
    brand = r["extracted"].get("brand") or ""
    heading = doc.add_heading(level=1)
    run = heading.add_run(f"{tier} · ")
    run.font.color.rgb = TIER_COLOR.get(tier, RGBColor(0, 0, 0))
    heading.add_run(campaign)
    if brand:
        sub = heading.add_run(f"  ·  {brand}")
        sub.font.size = Pt(13)
        sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Image
    try:
        img_buf = resize_for_doc(ROOT / r["file_path"])
        doc.add_picture(img_buf, width=Inches(6.5))
    except Exception as exc:  # noqa: BLE001
        p = doc.add_paragraph()
        p.add_run(f"(image not available: {exc})").italic = True

    # File ref (small, gray)
    file_p = doc.add_paragraph()
    file_run = file_p.add_run(r["file_path"])
    file_run.font.size = Pt(8)
    file_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # One-liner
    one_liner_p = doc.add_paragraph()
    one_liner_p.paragraph_format.space_before = Pt(8)
    run = one_liner_p.add_run(r["inferred"].get("one_liner") or "")
    run.bold = True
    run.font.size = Pt(15)

    # Score table
    table = doc.add_table(rows=2, cols=4)
    table.autofit = True
    headers = ["Idea  (20%)", "Strategy  (30%)", "Execution  (20%)", "Impact  (30%)"]
    scores = [w["idea"]["score"], w["strategy"]["score"], w["execution"]["score"], w["impact"]["score"]]
    for c, h in enumerate(headers):
        cell = table.rows[0].cells[c]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    for c, s in enumerate(scores):
        cell = table.rows[1].cells[c]
        cell.text = ""
        run = cell.paragraphs[0].add_run(f"{s}/100")
        run.bold = True
        run.font.size = Pt(14)

    # Weighted total + tier consistency
    total_p = doc.add_paragraph()
    total_p.paragraph_format.space_before = Pt(6)
    add_label(total_p, "Weighted total  ", size=11)
    total_run = total_p.add_run(f"{w['weighted_score']}/100  ")
    total_run.bold = True
    total_run.font.size = Pt(14)
    cons_run = total_p.add_run(f"({tier} band → {w['tier_consistency'].upper()})")
    cons_run.bold = True
    cons_run.font.size = Pt(10)
    cons_run.font.color.rgb = CONSISTENCY_COLOR.get(w["tier_consistency"], RGBColor(0, 0, 0))

    # Verdict (italic, indented vibe via a colored bullet)
    verdict_p = doc.add_paragraph()
    verdict_p.paragraph_format.space_before = Pt(10)
    add_label(verdict_p, "▎ JURY VERDICT  ", color=RGBColor(0xB8, 0x86, 0x0B), size=10)
    verdict_run = verdict_p.add_run(w["verdict"])
    verdict_run.italic = True
    verdict_run.font.size = Pt(11)

    # Per-axis rationale
    rat_heading = doc.add_paragraph()
    rat_heading.paragraph_format.space_before = Pt(12)
    add_label(rat_heading, "PER-AXIS RATIONALE", color=RGBColor(0x66, 0x66, 0x66), size=9)

    for k, label in [
        ("idea", "Idea"),
        ("strategy", "Strategy"),
        ("execution", "Execution"),
        ("impact", "Impact & Results"),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}  ·  {w[k]['score']}/100  ·  ")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(w[k]["rationale"]).font.size = Pt(10)

    # Creative mechanisms
    mech_p = doc.add_paragraph()
    mech_p.paragraph_format.space_before = Pt(8)
    add_label(mech_p, "Creative mechanisms  ", color=RGBColor(0x66, 0x66, 0x66), size=9)
    mechs = " · ".join(r["inferred"].get("creative_mechanisms") or [])
    mech_run = mech_p.add_run(mechs)
    mech_run.font.size = Pt(9)
    mech_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Feedback placeholder
    fb_p = doc.add_paragraph()
    fb_p.paragraph_format.space_before = Pt(14)
    add_label(fb_p, "→ FEEDBACK DA  ", color=RGBColor(0xB8, 0x86, 0x0B), size=10)
    hint = fb_p.add_run(
        "(highlight any text above and use Insert ▸ Comment, or Cmd + Option + M)"
    )
    hint.italic = True
    hint.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    hint.font.size = Pt(9)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, type=Path, help="Path to rationale JSONL (e.g. data/outdoor_2024_with_rationale.jsonl)")
    parser.add_argument("--output", required=True, type=Path, help="Path to write the .docx")
    parser.add_argument("--title", default="Mrs IArma · Review", help="Title of the doc")
    parser.add_argument("--subtitle", default="Cannes Lions winning boards · AI extraction + simulated jury rationale", help="Subtitle")
    args = parser.parse_args()

    in_path: Path = args.input if args.input.is_absolute() else ROOT / args.input
    out_path: Path = args.output if args.output.is_absolute() else ROOT / args.output

    records: list[dict] = []
    for line in in_path.open():
        line = line.strip()
        if not line:
            continue
        try:
            r = json.loads(line)
        except json.JSONDecodeError:
            continue
        if not r.get("error") and r.get("why_it_won"):
            records.append(r)

    records.sort(
        key=lambda r: (
            TIER_ORDER.get(r["tier"], 99),
            -r["why_it_won"]["weighted_score"],
        )
    )

    doc = Document()

    # Slightly tighter default font
    style = doc.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(11)

    # Title page
    title = doc.add_heading(args.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(args.subtitle)
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    instructions = doc.add_paragraph()
    instructions.alignment = WD_ALIGN_PARAGRAPH.CENTER
    instructions.paragraph_format.space_before = Pt(40)
    instr_run = instructions.add_run(
        "To leave feedback: highlight any text and use Insert ▸ Comment (Cmd + Opt + M).\n"
        "Comments are visible in real-time to the project team."
    )
    instr_run.font.size = Pt(10)
    instr_run.italic = True
    instr_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # TOC-like summary
    toc_heading = doc.add_paragraph()
    toc_heading.paragraph_format.space_before = Pt(40)
    add_label(toc_heading, "Contents", color=RGBColor(0x33, 0x33, 0x33), size=12)

    tier_counts: dict[str, int] = {}
    for r in records:
        tier_counts[r["tier"]] = tier_counts.get(r["tier"], 0) + 1
    for tier in ["Grand Prix", "Gold", "Silver", "Bronze", "Shortlist", "Loser"]:
        if tier_counts.get(tier):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            run = p.add_run(f"{tier}: ")
            run.bold = True
            run.font.color.rgb = TIER_COLOR[tier]
            p.add_run(f"{tier_counts[tier]} boards")

    # Boards
    for i, r in enumerate(records):
        doc.add_page_break()
        render_board(doc, r)

    out_path.parent.mkdir(exist_ok=True)
    doc.save(out_path)
    size_mb = out_path.stat().st_size / (1024 * 1024)
    print(f"Wrote {out_path}  ({len(records)} boards, {size_mb:.1f} MB)")


if __name__ == "__main__":
    main()
